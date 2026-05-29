"""Generate Security & Privacy Review document for PCF Workbench.

Aligns with Microsoft Secure Future Initiative (SFI), Security Development
Lifecycle (SDL), Secure by Design / Secure by Default principles, and the
Microsoft Privacy Standard. Output: docs/security-privacy-review.docx
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


REPO_ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = REPO_ROOT / "docs" / "security-privacy-review.docx"

ACCENT = RGBColor(0x00, 0x78, 0xD4)
MUTED = RGBColor(0x60, 0x60, 0x60)
BORDER = "BFBFBF"


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_table_borders(table, color: str = BORDER, size: str = "4") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        borders.append(b)
    tbl_pr.append(borders)


def style_heading(p, size_pt: int, color: RGBColor = ACCENT, bold: bool = True) -> None:
    for run in p.runs:
        run.font.size = Pt(size_pt)
        run.font.color.rgb = color
        run.font.bold = bold
        run.font.name = "Segoe UI"


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    sizes = {0: 28, 1: 18, 2: 14, 3: 12}
    style_heading(p, sizes.get(level, 12))


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False, color=None, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Segoe UI"
        run.font.size = Pt(11)


def add_kv_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for i, (k, v) in enumerate(rows):
        kc, vc = table.rows[i].cells
        kc.width = Cm(5.5)
        vc.width = Cm(11)
        kp = kc.paragraphs[0]
        kr = kp.add_run(k)
        kr.bold = True
        kr.font.name = "Segoe UI"
        kr.font.size = Pt(10)
        vp = vc.paragraphs[0]
        vr = vp.add_run(v)
        vr.font.name = "Segoe UI"
        vr.font.size = Pt(10)
        set_cell_shading(kc, "F2F6FB")
    set_table_borders(table)


def add_matrix_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        set_cell_shading(c, "0078D4")
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = table.rows[ri].cells[ci]
            p = c.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Segoe UI"
            r.font.size = Pt(10)
            if ri % 2 == 0:
                set_cell_shading(c, "F7F9FC")
    set_table_borders(table)


def build_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = title.add_run("Security & Privacy Review")
    tr.font.name = "Segoe UI Semibold"
    tr.font.size = Pt(32)
    tr.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sr = sub.add_run("PCF Workbench")
    sr.font.name = "Segoe UI"
    sr.font.size = Pt(20)
    sr.font.color.rgb = RGBColor(0x20, 0x20, 0x20)

    doc.add_paragraph()
    add_kv_table(
        doc,
        [
            ("Component", "PCF Workbench (open-source developer tool)"),
            ("Repository", "https://github.com/jaduplesms/PCF-Workbench (proposed transfer to microsoft/* org)"),
            ("Version under review", "1.0.0"),
            ("Document version", "1.0"),
            ("Date", date.today().isoformat()),
            ("Owner", "<TO BE FILLED - primary maintainer / FTE sponsor>"),
            ("Reviewers", "<TO BE FILLED - Security, Privacy, Legal/OSS>"),
            ("Classification", "Microsoft Confidential - Pre-publication"),
            ("Aligned standards", "Microsoft SFI, SDL, Secure by Design / Default, "
                                  "Microsoft Privacy Standard, OSSF Scorecard, OWASP ASVS L1"),
        ],
    )

    doc.add_paragraph()
    add_para(
        doc,
        "This document captures the security and privacy posture of PCF Workbench in support of "
        "publication under a Microsoft-owned GitHub organization. It follows the Microsoft "
        "Secure Future Initiative (SFI) pillars - Secure by Design, Secure by Default, Secure "
        "Operations - and the Microsoft Privacy Standard. Sections marked <TO BE FILLED> require "
        "input from the publishing team prior to sign-off.",
        italic=True,
        color=MUTED,
    )
    doc.add_page_break()


def build_overview(doc: Document) -> None:
    add_heading(doc, "1. Component overview", 1)
    add_para(
        doc,
        "PCF Workbench is a local development harness that replaces the built-in "
        "pcf-scripts start command for Power Apps Component Framework (PCF) controls. It "
        "loads a developer's compiled control bundle (out/controls/<Name>/bundle.js) into a "
        "browser page served by a local Vite dev server, shims the ComponentFramework.Context "
        "and Xrm.* namespaces, and provides developer-experience tooling: device emulation, "
        "network conditioning, in-memory WebAPI mocking from data.json, scenarios, lifecycle "
        "and resource-leak tracking, and a headless build-render-report loop for automated "
        "iteration.",
    )

    add_heading(doc, "1.1 Audience and use", 2)
    add_bullets(doc, [
        "Power Apps / Dynamics 365 developers building PCF controls.",
        "Runs only on the developer's workstation. There is no hosted service, no multi-tenant "
        "component, and no Microsoft-operated endpoint.",
        "Distribution: open source (MIT) via GitHub and npm. No installer, no auto-update.",
    ])

    add_heading(doc, "1.2 What it is NOT", 2)
    add_bullets(doc, [
        "Not a production runtime. It does not replace the Unified Client Interface - it emulates it for development.",
        "Not a service. There are no Microsoft-operated servers, no SaaS backend.",
        "Not a data store. It holds developer-supplied mock data in memory for the lifetime of the dev server process.",
        "Not bundled with Microsoft SDK source. All shims are original implementations; only types from "
        "@types/xrm and @types/powerapps-component-framework are referenced.",
    ])


def build_architecture(doc: Document) -> None:
    add_heading(doc, "2. Architecture and data flow", 1)
    add_para(doc, "The harness has two halves connected by a single Vite plugin:")
    add_bullets(doc, [
        "Server side (Node.js / Vite plugin): discovers PCF projects on disk, parses "
        "ControlManifest.Input.xml, serves the compiled bundle.js and bundled resources, "
        "watches out/ for HMR.",
        "Client side (React 18 + Zustand + Fluent UI v9): hosts the PCF control directly in "
        "the page (no iframe), provides side panels (Properties, Device, Network, WebAPI, "
        "Scenarios, Lifecycle, Renders, Form), and applies the shimmed Context to the control.",
        "Optional Dataverse bridge (vite-plugin/dataverse-proxy.ts, opt-in feature in roadmap M2): "
        "uses Microsoft Authentication Library (@azure/msal-node) with the Power Platform CLI "
        "(PAC) auth profile cache to mint access tokens server-side and proxy Web API calls. "
        "Tokens never leave the Node process.",
    ])

    add_heading(doc, "2.1 Trust boundaries", 2)
    add_matrix_table(
        doc,
        ["#", "Boundary", "Crosses"],
        [
            ["TB1", "Developer file system to Node Vite process",
             "Read-only access to PCF project files (manifest, bundle, data.json, scenarios, "
             "thumbnails) under user-specified paths."],
            ["TB2", "Node Vite process to Browser (localhost only)",
             "Static assets, manifest JSON, control bundle, harness UI. Default bind 127.0.0.1."],
            ["TB3", "Browser to user-supplied control code",
             "Control bundle executes in the harness page origin with full DOM access (same model "
             "as pcf-scripts start). The harness monkey-patches listener/timer/observer APIs "
             "for leak detection only."],
            ["TB4", "Node process to PAC CLI cache (optional, Dataverse Bridge)",
             "Reads %LOCALAPPDATA%\\Microsoft\\PowerAppsCli\\authprofiles_v2.json and the "
             "DPAPI-protected MSAL cache. No write-back."],
            ["TB5", "Node process to Dataverse org HTTPS endpoint (optional, Dataverse Bridge)",
             "Outbound HTTPS to the org URL recorded in the active PAC profile. Read-only by "
             "default; writes require explicit per-call confirmation in the UI."],
        ],
    )

    add_heading(doc, "2.2 Inputs / outputs", 2)
    add_matrix_table(
        doc,
        ["Surface", "Source", "Sink", "Notes"],
        [
            ["ControlManifest.Input.xml", "Disk (user PCF project)", "Server parser to JSON to browser",
             "Parsed with fast-xml-parser; not executed."],
            ["bundle.js", "Disk (out/controls/<Name>/)", "Browser <script>", "Executed in harness origin. "
             "Same trust as developer's own code."],
            ["data.json", "Disk", "In-memory store (browser + server cache)",
             "Mock entity rows. Treated as developer-supplied test data."],
            ["test-scenarios.json", "Disk", "In-memory store", "User-saved property + viewport snapshots."],
            ["thumbnail.{jpg,png,gif}", "Disk", "Browser <img>", "Static asset."],
            ["Dataverse Web API response (optional)", "External HTTPS", "Browser (proxied)",
             "Cached on disk under user-controlled directory; cache is opt-in."],
            ["Loop report.json + screenshot.png", "Headless run", "User-specified --out dir",
             "Created by the developer's own command; no upload."],
        ],
    )


def build_secure_by_design(doc: Document) -> None:
    add_heading(doc, "3. Secure by Design", 1)
    add_para(doc, "How design-time decisions reduce attack surface and blast radius.")

    add_heading(doc, "3.1 Principle alignment (SFI / SDL)", 2)
    add_matrix_table(
        doc,
        ["SFI / SDL principle", "How PCF Workbench applies it"],
        [
            ["Protect identities and secrets",
             "No identities are managed by the harness. The optional Dataverse Bridge delegates entirely "
             "to PAC CLI / MSAL; the harness never prompts for passwords, never stores credentials, and "
             "never exposes access tokens to the browser. Tokens are held in process memory only and "
             "rely on MSAL's DPAPI-protected persistence."],
            ["Protect tenants and isolate production systems",
             "Default operating mode uses mock data only (data.json). The live Dataverse bridge is "
             "off by default, read-only by default when enabled, and writes require explicit per-call "
             "user confirmation in the UI. Network mode 'offline' simulates disconnection without any "
             "external calls."],
            ["Protect networks",
             "Vite dev server binds to localhost. Documentation directs operators to use --host "
             "127.0.0.1 for deterministic local-only binding. No inbound public endpoints. No "
             "TLS-terminating server - the harness never exposes a public surface."],
            ["Protect engineering systems (supply chain)",
             "Dependencies pinned in package.json with explicit version ranges; package-lock.json "
             "committed. CI runs npm audit and npm run typecheck on every PR. Dependabot enabled. "
             "GitHub branch protection requires reviews. SBOM produced via npm sbom at release."],
            ["Monitor and detect threats",
             "All shim activity (WebAPI calls, lifecycle events, listener/timer/observer creation) is "
             "captured in the harness UI. Resource leaks are diff-reported on destroy. Errors are "
             "surfaced via the dialog bus and console. No telemetry leaves the machine."],
            ["Accelerate response and remediation",
             "Repository will adopt SECURITY.md with private vulnerability reporting and a defined "
             "SLA for triage / fix / advisory. CodeQL enabled on default branch."],
        ],
    )

    add_heading(doc, "3.2 Secure defaults", 2)
    add_bullets(doc, [
        "Dataverse Bridge is OFF by default. Mock data (data.json) is the only data source out of the box.",
        "Live Dataverse access, when enabled, is READ-ONLY by default; writes require explicit per-call confirmation.",
        "Vite dev server runs on localhost; documentation steers users to bind 127.0.0.1 explicitly.",
        "No external <script src=...> tags are injected by the harness; the user's compiled bundle is the only "
        "untrusted code executed in the page (this matches pcf-scripts start).",
        "Control CSS is sandboxed into @layer pcf-control and @media rules are rewritten to "
        "@container pcf-viewport, preventing accidental host-style takeover.",
        "Strict TypeScript mode enforced (tsc --noEmit is the pre-PR gate).",
    ])

    add_heading(doc, "3.3 Threat model (STRIDE summary)", 2)
    add_matrix_table(
        doc,
        ["Threat", "Asset / surface", "Impact", "Mitigation", "Residual risk"],
        [
            ["Spoofing", "Dataverse Bridge identity",
             "Caller authenticates as wrong user.",
             "Identity is delegated to MSAL + PAC CLI cache. The harness does not perform its own "
             "auth. Active profile is shown in the UI before any request.", "Low"],
            ["Tampering", "User-supplied bundle.js",
             "Malicious developer-authored code runs in harness origin.",
             "Same trust model as pcf-scripts start - the developer owns the code. The harness is "
             "a dev tool, not a sandbox. CSS-layer isolation prevents accidental UI takeover.",
             "Accepted - design constraint of local dev tooling"],
            ["Repudiation", "Local activity",
             "Loss of audit trail.",
             "All Dataverse Bridge requests are logged to the in-memory WebAPI log with timestamps.",
             "Low (local dev only)"],
            ["Information disclosure", "Access tokens",
             "Token exfiltration could compromise Dataverse org.",
             "Tokens are server-side only. The proxy strips authorization headers from "
             "responses returned to the browser. The PublicProfile shape explicitly excludes "
             "sensitive PAC profile fields (object IDs, raw tokens).", "Low"],
            ["Information disclosure", "Mock data (data.json)",
             "Developer commits real customer data into mock store.",
             "Documented guidance: mock data must be synthetic. CONTRIBUTING.md will require "
             "samples to use generated/fake data only. Pre-commit advisory will note this.",
             "Accepted with guidance"],
            ["Denial of service", "Vite dev server",
             "Resource exhaustion crashes dev loop.",
             "Local-only, single-user. Not in scope for hardening beyond Vite defaults.",
             "Accepted"],
            ["Elevation of privilege", "Loaded control",
             "Control reads files outside its directory.",
             "Browser sandbox limits file access. The Vite plugin only serves files under the "
             "configured workspace / control root.", "Low"],
            ["Supply chain", "npm dependencies",
             "Compromised dependency leaks dev environment data.",
             "Pinned lockfile, Dependabot, npm audit in CI, CodeQL. Plan to enable OSSF "
             "Scorecard reporting and OIDC-based provenance for npm publishes.", "Low-Medium"],
            ["Cache poisoning", "Dataverse response cache",
             "Stale or attacker-supplied data persists across runs.",
             "Cache is opt-in, scoped under user-controlled directory, and is read on a same-user "
             "basis only. Documentation directs users to clear cache when switching orgs.",
             "Low"],
        ],
    )


def build_secure_by_default(doc: Document) -> None:
    add_heading(doc, "4. Secure by Default", 1)
    add_para(doc, "What the user gets in the box without configuration.")
    add_matrix_table(
        doc,
        ["Behavior", "Default", "Why it is the safe default"],
        [
            ["Bind address", "127.0.0.1", "Prevents accidental LAN exposure."],
            ["External network calls", "None", "Mock WebAPI is in-memory only."],
            ["Dataverse Bridge", "Disabled", "Opt-in only after explicit env-var / config."],
            ["Live Dataverse writes", "Disabled (read-only)", "Avoids accidental data mutation against real orgs."],
            ["Telemetry", "None collected", "No analytics endpoint exists."],
            ["Persistent storage of user data", "None", "All mock data is in-memory; user files are read-only."],
            ["Auto-update", "Disabled", "Distribution is via developer-controlled npm/git."],
            ["Logs containing sensitive data", "Suppressed", "Authorization headers redacted from request log."],
        ],
    )


def build_privacy(doc: Document) -> None:
    add_heading(doc, "5. Privacy assessment (Microsoft Privacy Standard)", 1)

    add_heading(doc, "5.1 Personal data collected by the product", 2)
    add_para(doc, "None. PCF Workbench does not collect, transmit, or store personal data on behalf of Microsoft.")

    add_heading(doc, "5.2 Personal data observed during operation", 2)
    add_bullets(doc, [
        "When the optional Dataverse Bridge is enabled, the harness will pass through whatever "
        "the connected Dataverse org returns. If that data includes personal data, it is the "
        "developer's responsibility under the Microsoft Privacy Standard to operate against an "
        "appropriately-scoped non-production environment.",
        "PAC profile metadata (signed-in user UPN, tenant id, environment URL) is read from the "
        "PAC CLI cache to display the active connection. This metadata is shown to the user only "
        "and never transmitted off-machine.",
        "Mock data in data.json is supplied by the developer and is out of scope for the product.",
    ])

    add_heading(doc, "5.3 Telemetry and analytics", 2)
    add_matrix_table(
        doc,
        ["Category", "Status"],
        [
            ["Required service data", "None"],
            ["Optional diagnostic data", "None"],
            ["Crash dumps", "None uploaded; errors stay in local console / UI"],
            ["Usage analytics", "None"],
            ["Third-party telemetry SDKs", "None (verified absent from package.json)"],
        ],
    )

    add_heading(doc, "5.4 Data residency / cross-border transfer", 2)
    add_para(
        doc,
        "All data stays on the developer's workstation. The optional Dataverse Bridge connects "
        "to whatever org the developer's PAC profile already points at - residency is governed by "
        "that environment's geography, not by this tool.",
    )

    add_heading(doc, "5.5 GDPR / DSR considerations", 2)
    add_bullets(doc, [
        "Right of access / portability: no Microsoft-held data exists to disclose.",
        "Right to erasure: developer can delete the repo / npm cache / local config directories at any time.",
        "Sub-processors: none.",
        "Cookies / tracking technologies: none.",
    ])

    add_heading(doc, "5.6 Accessibility and inclusive design", 2)
    add_bullets(doc, [
        "Harness UI built on Fluent UI v9, which ships WCAG 2.1 AA-aligned components.",
        "High-contrast and dark-mode preview toggles planned in roadmap M8.",
        "Roadmap M4 (Diagnostics and Linting) integrates axe-core a11y audits against user controls.",
    ])


def build_secrets_and_creds(doc: Document) -> None:
    add_heading(doc, "6. Secrets, credentials, and cryptography", 1)
    add_bullets(doc, [
        "No secrets are committed to the repository. Pre-commit and CI gitleaks scan planned.",
        "No symmetric / asymmetric keys are generated or managed by the product.",
        "Optional Dataverse Bridge uses MSAL public client flow with PKCE; refresh tokens are "
        "persisted by MSAL using OS-provided protection (DPAPI on Windows, Keychain on macOS, "
        "libsecret on Linux). The harness never reads raw tokens off disk and never writes them.",
        "All outbound calls (when Dataverse Bridge is enabled) use HTTPS via node:https. "
        "Certificate validation uses Node defaults; no pinning override.",
        "No custom cryptography implemented. No cryptographic primitives shipped.",
    ])


def build_dependencies(doc: Document) -> None:
    add_heading(doc, "7. Supply chain and dependencies", 1)
    add_matrix_table(
        doc,
        ["Layer", "Notable dependencies", "Notes"],
        [
            ["Runtime UI", "react, react-dom, @fluentui/react-components, @fluentui/react-icons, zustand",
             "Microsoft-published Fluent UI; React 18 LTS."],
            ["Auth (optional)", "@azure/msal-node, @azure/msal-node-extensions",
             "Microsoft-published MSAL libraries."],
            ["Dev server", "vite, @vitejs/plugin-react", "Local-only dev server."],
            ["Tooling", "typescript, tsx, commander, fast-xml-parser, html2canvas, @playwright/test",
             "Standard ecosystem tools."],
            ["Type-only", "@types/xrm, @types/powerapps-component-framework, @types/node, @types/react",
             "Types only; no runtime code."],
        ],
    )
    add_bullets(doc, [
        "Lockfile committed (package-lock.json). Reproducible installs via npm ci.",
        "Dependabot to be enabled for the GitHub repository before publication.",
        "GitHub Advanced Security (CodeQL, secret scanning, dependency review) to be enabled on transfer.",
        "Release process: provenance-attested npm publish via OIDC; tagged GitHub release with SBOM (CycloneDX).",
        "No bundled forks of Microsoft SDK source. PCF and Xrm types are referenced only, not copied.",
    ])


def build_operations(doc: Document) -> None:
    add_heading(doc, "8. Secure Operations", 1)
    add_bullets(doc, [
        "Vulnerability reporting: SECURITY.md in the repository directs reports to GitHub Private "
        "Vulnerability Reporting (GHSA). SLA: triage within 5 business days; fix within 30 days for "
        "Critical/High, 90 days for Medium/Low.",
        "Branch protection: required reviews, status checks, signed commits where feasible.",
        "Required CI checks: npm run typecheck, npm audit --omit=dev (non-blocking advisory), "
        "Playwright conformance smoke against the ConformanceTester sample, CodeQL.",
        "Release cadence: semantic versioning; security patches released as patch versions and "
        "documented in CHANGELOG.md with a GHSA advisory.",
        "Incident response: maintainer + Microsoft sponsor on-call list maintained privately. "
        "Coordinated disclosure via GitHub Security Advisories.",
        "End of life: deprecation notice committed to repository at least 90 days before archive.",
    ])


def build_compliance(doc: Document) -> None:
    add_heading(doc, "9. Compliance and legal", 1)
    add_kv_table(
        doc,
        [
            ("License", "MIT (proposed to remain MIT on transfer; subject to Microsoft OSS review)"),
            ("Third-party notices", "THIRD-PARTY-NOTICES.md to be generated at release"),
            ("Cryptography export classification", "Not applicable - no proprietary cryptography shipped"),
            ("Trademark", "Adopts Microsoft Trademark and Brand Guidelines on transfer"),
            ("Code of Conduct", "Microsoft Open Source Code of Conduct"),
            ("Contributor License Agreement", "Microsoft CLA via cla.opensource.microsoft.com"),
            ("Accessibility conformance", "Targets WCAG 2.1 AA; statement to be published with first Microsoft release"),
            ("Data Protection Impact Assessment", "Not required (no personal data processed)"),
        ],
    )


def build_open_items(doc: Document) -> None:
    add_heading(doc, "10. Open items prior to publication", 1)
    add_bullets(doc, [
        "Add SECURITY.md with private vulnerability reporting contact.",
        "Add CODE_OF_CONDUCT.md (Microsoft template).",
        "Add SUPPORT.md clarifying community-only support.",
        "Add .github/dependabot.yml and enable Dependabot security updates.",
        "Enable GitHub Advanced Security (CodeQL, secret scanning, push protection).",
        "Adopt the Microsoft OSPO repository template / checks.",
        "Generate SBOM (CycloneDX) as part of the release workflow.",
        "Update README.md with explicit safety guidance for the Dataverse Bridge.",
        "Run a final dependency audit and capture results as an appendix to this document.",
        "Sign off: Security reviewer, Privacy reviewer, OSPO, FTE sponsor.",
    ])


def build_signoff(doc: Document) -> None:
    add_heading(doc, "11. Approval and sign-off", 1)
    add_matrix_table(
        doc,
        ["Role", "Name", "Date", "Decision"],
        [
            ["FTE sponsor", "", "", ""],
            ["Security reviewer", "", "", ""],
            ["Privacy reviewer", "", "", ""],
            ["Open Source Programs Office (OSPO)", "", "", ""],
            ["Accessibility reviewer (optional)", "", "", ""],
            ["Engineering lead", "", "", ""],
        ],
    )
    doc.add_paragraph()
    add_para(
        doc,
        "Sign-off indicates that the reviewer has read this document, confirmed the controls "
        "described are implemented in the version under review, and has no blocking concerns. "
        "Any non-blocking follow-ups must be tracked as GitHub issues with owners before merge.",
        italic=True, color=MUTED,
    )


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(11)

    build_cover(doc)
    build_overview(doc)
    build_architecture(doc)
    build_secure_by_design(doc)
    build_secure_by_default(doc)
    build_privacy(doc)
    build_secrets_and_creds(doc)
    build_dependencies(doc)
    build_operations(doc)
    build_compliance(doc)
    build_open_items(doc)
    build_signoff(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
